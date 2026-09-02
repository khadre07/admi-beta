"""Import / export Excel pour ADMI.

Reconnaît automatiquement les feuilles « Machines », « Arrêts », « Énergie »,
« Interventions » et « Planning » et fait correspondre les colonnes de façon
souple (accents, casse, libellés approchants), comme l'application d'origine.
"""
from __future__ import annotations

import io
import unicodedata
from datetime import date, datetime

import pandas as pd

from .config import (DEPARTEMENTS, MOIS, MOIS_COURT, STATUTS_MACHINE,
                     TYPES_ARRET, TYPES_INTERV, TYPES_PLAN, dep)
from .data import Database, uid
from .kpis import hours_between, intervention_cost
from .stock import piece_valeur

TYPES = ["machines", "arrets", "energie", "interventions", "planning"]
LABELS = {"machines": "Machines", "arrets": "Arrêts", "energie": "Énergie",
          "interventions": "Interventions", "planning": "Planning"}

SHEET_MATCHERS = {
    "machines": ["machine", "equipement", "parc"],
    "arrets": ["arret", "downtime"],
    "energie": ["energie", "electricite"],
    "interventions": ["intervention", "rapport"],
    "planning": ["planning", "plan", "calendrier"],
}


# ---------------------------------------------------------------------------
# Normalisation & extraction souple
# ---------------------------------------------------------------------------
def normalize_str(s) -> str:
    if s is None:
        return ""
    s = str(s)
    s = unicodedata.normalize("NFD", s)
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    return s.lower().strip()


def get_val(row: dict, keywords):
    for kw in keywords:
        nkw = normalize_str(kw)
        for key, val in row.items():
            if nkw in normalize_str(key) and str(val).strip() != "":
                return val
    return ""


def match_dept_flexible(value):
    n = normalize_str(value)
    if not n:
        return None
    for d in DEPARTEMENTS:
        if n in (normalize_str(d["nom"]), normalize_str(d["court"]), normalize_str(d["id"])):
            return d
    for d in DEPARTEMENTS:
        if n in normalize_str(d["nom"]) or normalize_str(d["court"]) in n:
            return d
    return None


def _to_dt(value):
    if isinstance(value, (datetime, pd.Timestamp)):
        return value.to_pydatetime() if isinstance(value, pd.Timestamp) else value
    if isinstance(value, date):
        return datetime(value.year, value.month, value.day)
    return None


def parse_date_flexible(value):
    dt = _to_dt(value)
    if dt:
        return dt.strftime("%Y-%m-%d")
    s = str(value or "").strip()
    if not s:
        return None
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y/%m/%d"):
        try:
            return datetime.strptime(s[:10], fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return None


def parse_datetime_flexible(value):
    dt = _to_dt(value)
    if dt:
        return dt.strftime("%Y-%m-%dT%H:%M")
    s = str(value or "").strip()
    if not s:
        return None
    for fmt in ("%Y-%m-%d %H:%M", "%Y-%m-%dT%H:%M", "%d/%m/%Y %H:%M", "%Y-%m-%d %H:%M:%S"):
        try:
            return datetime.strptime(s, fmt).strftime("%Y-%m-%dT%H:%M")
        except ValueError:
            continue
    d = parse_date_flexible(value)
    return d + "T00:00" if d else None


def parse_mois(value):
    """Accepte un numéro (1-12) ou un nom de mois (« Janvier », « janv »…).
    Retourne l'index 0-11, ou None si non reconnu."""
    s = str(value or "").strip()
    if not s:
        return None
    try:
        n = int(float(s))
        if 1 <= n <= 12:
            return n - 1
    except ValueError:
        pass
    ns = normalize_str(s)
    for i, m in enumerate(MOIS):
        if normalize_str(m) == ns or normalize_str(MOIS_COURT[i]) == ns or ns.startswith(normalize_str(m)[:4]):
            return i
    return None


def match_machine(name, db: Database, pending):
    n = normalize_str(name)
    if not n:
        return None
    allm = list(db.machines) + list(pending)
    for m in allm:
        if normalize_str(m["nom"]) == n:
            return m
    for m in allm:
        if n in normalize_str(m["nom"]) or normalize_str(m["nom"]) in n:
            return m
    return None


def resolve_or_create_machine(name, dept_raw, db, pending, errors, sheet, line):
    m = match_machine(name, db, pending)
    if m:
        return m["id"]
    d = match_dept_flexible(dept_raw)
    if not d:
        return None  # machine inconnue et pas de département -> ligne rejetée
    new = {"id": uid(), "nom": str(name).strip(), "departementId": d["id"],
           "puissanceKW": 0, "dateMES": "", "statut": "En service", "autoCree": True}
    pending.append(new)
    return new["id"]


# ---------------------------------------------------------------------------
# Ingestion d'une feuille (liste de dictionnaires) -> enregistrements ADMI
# ---------------------------------------------------------------------------
def _parse_pieces(raw):
    pieces = []
    for chunk in str(raw or "").split(";"):
        chunk = chunk.strip()
        if not chunk:
            continue
        parts = chunk.split(":")
        desig = parts[0].strip()
        if not desig:
            continue
        qte = float(parts[1]) if len(parts) > 1 and parts[1].strip() else 1
        cout = float(parts[2]) if len(parts) > 2 and parts[2].strip() else 0
        pieces.append({"designation": desig, "qte": qte, "cout": cout})
    return pieces


def ingest_rows(typ, rows, sheet, result, db):
    pending_machines = result["machines"]
    for idx, row in enumerate(rows, start=2):  # ligne 1 = en-tête
        try:
            if typ == "machines":
                nom = get_val(row, ["nom", "machine", "equipement"])
                if not str(nom).strip():
                    continue
                d = match_dept_flexible(get_val(row, ["departement", "service"]))
                if not d:
                    result["errors"].append({"type": typ, "sheet": sheet, "line": idx,
                                             "message": "département manquant ou non reconnu"})
                    continue
                pending_machines.append({
                    "id": uid(), "nom": str(nom).strip(), "departementId": d["id"],
                    "puissanceKW": float(get_val(row, ["puissance", "kw"]) or 0),
                    "dateMES": parse_date_flexible(get_val(row, ["mise en service", "date"])) or "",
                    "statut": str(get_val(row, ["statut", "etat"]) or "En service").strip() or "En service",
                })

            elif typ == "arrets":
                mach = get_val(row, ["machine", "equipement"])
                mid = resolve_or_create_machine(mach, get_val(row, ["departement", "service"]),
                                                db, pending_machines, result["errors"], sheet, idx)
                debut = parse_datetime_flexible(get_val(row, ["debut", "start"]))
                fin = parse_datetime_flexible(get_val(row, ["fin", "end"]))
                if not mid or not debut or not fin:
                    result["errors"].append({"type": typ, "sheet": sheet, "line": idx,
                                             "message": "machine/début/fin requis (ajoutez un département si la machine est nouvelle)"})
                    continue
                m = match_machine(mach, db, pending_machines)
                result["arrets"].append({
                    "id": uid(), "machineId": mid, "departementId": m["departementId"],
                    "type": str(get_val(row, ["type"]) or "Panne").strip() or "Panne",
                    "cause": str(get_val(row, ["cause"]) or "").strip(),
                    "dateDebut": debut, "dateFin": fin,
                    "description": str(get_val(row, ["description"]) or "").strip(),
                })

            elif typ == "energie":
                d = match_dept_flexible(get_val(row, ["departement", "service"]))
                mois = parse_mois(get_val(row, ["mois"]))
                annee = get_val(row, ["annee", "year"])
                if not d or mois is None or str(annee) == "":
                    result["errors"].append({"type": typ, "sheet": sheet, "line": idx,
                                             "message": "département / mois / année requis (mois = 1-12 ou nom)"})
                    continue
                result["energie"].append({
                    "id": uid(), "departementId": d["id"], "mois": mois, "annee": int(float(annee)),
                    "kwh": float(get_val(row, ["consommation", "kwh"]) or 0),
                    "montant": float(get_val(row, ["montant", "cout", "fcfa"]) or 0),
                })

            elif typ == "interventions":
                mach = get_val(row, ["machine", "equipement"])
                mid = resolve_or_create_machine(mach, get_val(row, ["departement", "service"]),
                                                db, pending_machines, result["errors"], sheet, idx)
                idate = parse_date_flexible(get_val(row, ["date"]))
                if not mid or not idate:
                    result["errors"].append({"type": typ, "sheet": sheet, "line": idx,
                                             "message": "machine et date requises"})
                    continue
                m = match_machine(mach, db, pending_machines)
                result["interventions"].append({
                    "id": uid(), "machineId": mid, "departementId": m["departementId"], "date": idate,
                    "type": str(get_val(row, ["type"]) or "Correctif").strip() or "Correctif",
                    "technicien": str(get_val(row, ["technicien"]) or "").strip(),
                    "duree": float(get_val(row, ["duree"]) or 0),
                    "coutMainOeuvre": float(get_val(row, ["main d'oeuvre", "main d'œuvre", "mo"]) or 0),
                    "description": str(get_val(row, ["description"]) or "").strip(),
                    "pieces": _parse_pieces(get_val(row, ["piece"])),
                })

            elif typ == "planning":
                mach = get_val(row, ["machine", "equipement"])
                mid = resolve_or_create_machine(mach, get_val(row, ["departement", "service"]),
                                                db, pending_machines, result["errors"], sheet, idx)
                titre = get_val(row, ["titre"])
                pdate = parse_date_flexible(get_val(row, ["date"]))
                if not mid or not str(titre).strip() or not pdate:
                    result["errors"].append({"type": typ, "sheet": sheet, "line": idx,
                                             "message": "titre, machine et date requis"})
                    continue
                m = match_machine(mach, db, pending_machines)
                result["planning"].append({
                    "id": uid(), "machineId": mid, "departementId": m["departementId"],
                    "titre": str(titre).strip(), "date": pdate,
                    "type": str(get_val(row, ["type"]) or "Préventif").strip() or "Préventif",
                    "statut": str(get_val(row, ["statut"]) or "Planifié").strip() or "Planifié",
                    "description": str(get_val(row, ["description"]) or "").strip(),
                })
        except (ValueError, TypeError) as exc:
            result["errors"].append({"type": typ, "sheet": sheet, "line": idx, "message": str(exc)})


def detect_sheet_type(sheet_name, columns=None):
    n = normalize_str(sheet_name)
    for typ, kws in SHEET_MATCHERS.items():
        if any(k in n for k in kws):
            return typ
    # repli : détection par colonnes (utile pour CSV)
    if columns:
        cols = " ".join(normalize_str(c) for c in columns)
        if "puissance" in cols and "machine" in cols:
            return "machines"
        if "debut" in cols and "fin" in cols:
            return "arrets"
        if "kwh" in cols or "consommation" in cols:
            return "energie"
        if "technicien" in cols or "main d'oeuvre" in normalize_str(cols):
            return "interventions"
        if "titre" in cols:
            return "planning"
    return None


def parse_workbook(file_bytes: bytes, filename: str, db: Database) -> dict:
    result = {t: [] for t in TYPES}
    result["errors"] = []
    result["fileName"] = filename

    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "csv":
        df = pd.read_csv(io.BytesIO(file_bytes), dtype=object, keep_default_na=False)
        sheets = {filename: df}
    else:
        sheets = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None, dtype=object)

    # traiter les machines d'abord pour que les autres feuilles les retrouvent
    ordered = sorted(sheets.items(),
                     key=lambda kv: 0 if detect_sheet_type(kv[0], kv[1].columns) == "machines" else 1)
    for sheet_name, df in ordered:
        typ = detect_sheet_type(sheet_name, df.columns)
        if not typ:
            continue
        df = df.where(pd.notna(df), "")
        rows = df.to_dict(orient="records")
        ingest_rows(typ, rows, sheet_name, result, db)
    return result


# ---------------------------------------------------------------------------
# Import Word (.docx) et PDF — les données doivent être présentées en TABLEAUX
# avec une ligne d'en-tête (comme dans le modèle Excel).
# ---------------------------------------------------------------------------
def _rows_from_grid(grid):
    """(en-têtes + lignes) -> liste de dicts. La 1re ligne est l'en-tête."""
    headers = [str(h or "").strip() for h in grid[0]]
    rows = []
    for cells in grid[1:]:
        cells = [str(c or "").strip() for c in cells]
        rows.append({headers[j]: (cells[j] if j < len(cells) else "") for j in range(len(headers))})
    return headers, rows


def _ingest_tables(tables, filename, db):
    """tables : liste de dicts {headers, rows, label}. Renvoie un result ADMI."""
    result = {t: [] for t in TYPES}
    result["errors"] = []
    result["fileName"] = filename
    # traiter les tableaux « machines » d'abord (pour que les autres les retrouvent)
    ordered = sorted(tables, key=lambda t: 0 if detect_sheet_type(None, t["headers"]) == "machines" else 1)
    recognized = 0
    for t in ordered:
        typ = detect_sheet_type(None, t["headers"])
        if not typ:
            continue
        recognized += 1
        ingest_rows(typ, t["rows"], t["label"], result, db)
    if recognized == 0:
        raise ValueError(
            "Aucun tableau reconnaissable n'a été trouvé. Les données doivent être "
            "présentées en tableau avec une ligne d'en-tête (colonnes comme dans le "
            "modèle Excel). Pour un document en texte libre, utilisez plutôt le modèle Excel.")
    return result


def parse_docx(file_bytes: bytes, filename: str, db: Database) -> dict:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    tables = []
    for ti, table in enumerate(doc.tables, start=1):
        grid = [[cell.text for cell in row.cells] for row in table.rows]
        if len(grid) < 2:
            continue
        headers, rows = _rows_from_grid(grid)
        tables.append({"headers": headers, "rows": rows, "label": f"Tableau Word #{ti}"})
    return _ingest_tables(tables, filename, db)


def parse_pdf(file_bytes: bytes, filename: str, db: Database) -> dict:
    import pymupdf
    tables = []
    idx = 0
    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    try:
        for page in doc:
            for tab in page.find_tables().tables:
                grid = tab.extract()
                if not grid or len(grid) < 2:
                    continue
                idx += 1
                headers, rows = _rows_from_grid(grid)
                tables.append({"headers": headers, "rows": rows, "label": f"Tableau PDF #{idx}"})
    finally:
        doc.close()
    return _ingest_tables(tables, filename, db)


def parse_import(file_bytes: bytes, filename: str, db: Database) -> dict:
    """Point d'entrée unique : route selon l'extension du fichier."""
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext in ("xlsx", "xls", "csv"):
        return parse_workbook(file_bytes, filename, db)
    if ext == "docx":
        return parse_docx(file_bytes, filename, db)
    if ext == "pdf":
        return parse_pdf(file_bytes, filename, db)
    raise ValueError("Format non pris en charge (utilisez .xlsx, .xls, .csv, .docx ou .pdf).")


def apply_import(db: Database, result: dict, mode: str = "append") -> int:
    total = 0
    for typ in TYPES:
        rows = result.get(typ) or []
        if not rows:
            continue
        # nettoyer le marqueur autoCree
        clean = [{k: v for k, v in r.items() if k != "autoCree"} for r in rows]
        total += len(clean)
        if typ == "machines":
            if mode == "replace":
                db.machines = clean
            else:
                for newm in clean:
                    existing = next((m for m in db.machines
                                     if normalize_str(m["nom"]) == normalize_str(newm["nom"])), None)
                    if existing:
                        existing.update({k: v for k, v in newm.items() if k != "id"})
                    else:
                        db.machines.append(newm)
        else:
            target = getattr(db, typ)
            setattr(db, typ, clean if mode == "replace" else target + clean)
    return total


# ---------------------------------------------------------------------------
# Export & modèle
# ---------------------------------------------------------------------------
def _write_sheets(sheets: dict) -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as xl:
        for name, (header, rows) in sheets.items():
            pd.DataFrame(rows, columns=header).to_excel(xl, sheet_name=name, index=False)
    return buf.getvalue()


def export_bytes(db: Database) -> bytes:
    sheets = {
        "Machines": (["Nom machine", "Département", "Puissance (kW)", "Date mise en service", "Statut"],
                     [[m["nom"], dep(m["departementId"])["nom"], m["puissanceKW"], m["dateMES"], m["statut"]]
                      for m in db.machines]),
        "Arrêts": (["Machine", "Département", "Type d'arrêt", "Cause", "Début", "Fin", "Durée (h)", "Description"],
                   [[db.machine_name(a["machineId"]), dep(a["departementId"])["nom"], a["type"], a.get("cause", ""),
                     a["dateDebut"], a["dateFin"], round(hours_between(a["dateDebut"], a["dateFin"]), 2),
                     a.get("description", "")] for a in db.arrets]),
        "Énergie": (["Département / Service", "Mois", "Année", "Consommation (kWh)", "Montant (FCFA)"],
                    [[dep(e["departementId"])["nom"], MOIS[e["mois"]], e["annee"], e["kwh"], e["montant"]]
                     for e in db.energie]),
        "Interventions": (["Date", "Machine", "Département", "Type", "Technicien", "Durée (h)",
                           "Coût main d'œuvre", "Coût pièces", "Coût total", "Description"],
                          [[i["date"], db.machine_name(i["machineId"]), dep(i["departementId"])["nom"], i["type"],
                            i.get("technicien", ""), i.get("duree", 0), i.get("coutMainOeuvre", 0),
                            sum((p.get("cout", 0) * p.get("qte", 1)) for p in i.get("pieces", [])),
                            intervention_cost(i), i.get("description", "")] for i in db.interventions]),
        "Pièces": (["Désignation", "Référence", "Département", "Quantité", "Unité", "Seuil d'alerte",
                    "Coût unitaire (FCFA)", "Valeur (FCFA)", "Emplacement", "Fournisseur"],
                   [[p["designation"], p.get("reference", ""),
                     dep(p["departementId"])["nom"] if p.get("departementId") else "",
                     p.get("quantite", 0), p.get("unite", ""), p.get("seuilAlerte", 0),
                     p.get("coutUnitaire", 0), piece_valeur(p), p.get("emplacement", ""),
                     p.get("fournisseur", "")] for p in db.pieces]),
        "Mouvements": (["Date", "Pièce", "Type", "Quantité", "Motif"],
                       [[m.get("date", ""), _piece_name(db, m["pieceId"]), m["type"],
                         m.get("quantite", 0), m.get("motif", "")]
                        for m in sorted(db.mouvements, key=lambda x: x.get("date", ""))]),
    }
    return _write_sheets(sheets)


def _piece_name(db: Database, piece_id: str) -> str:
    return next((p["designation"] for p in db.pieces if p["id"] == piece_id), "")


def _piece_name(db: Database, piece_id: str) -> str:
    return next((p["designation"] for p in db.pieces if p["id"] == piece_id), "")


def template_bytes() -> bytes:
    sheets = {
        "Machines": (["Nom machine", "Département", "Puissance (kW)", "Date mise en service (AAAA-MM-JJ)", "Statut"],
                     [["Presse à injection AM-3", "Articles Ménagers", 45, "2023-03-01", "En service"],
                      ["Onduleuse Zinc ON-3", "Ondulations (Toitures Zinc)", 55, "2021-06-15", "En service"]]),
        "Arrêts": (["Machine", "Département (optionnel)", "Type d'arrêt", "Cause",
                    "Début (AAAA-MM-JJ HH:MM)", "Fin (AAAA-MM-JJ HH:MM)", "Description"],
                   [["Presse à injection AM-1", "Articles Ménagers", "Panne", "Rupture courroie",
                     "2026-01-10 08:00", "2026-01-10 11:30", "Remplacement de la courroie"]]),
        "Énergie": (["Département / Service", "Mois (1-12)", "Année", "Consommation (kWh)", "Montant (FCFA)"],
                    [["Articles Ménagers", 1, 2026, 15000, 2250000]]),
        "Interventions": (["Date (AAAA-MM-JJ)", "Machine", "Département (optionnel)", "Type intervention",
                           "Technicien", "Durée (h)", "Coût main d'œuvre (FCFA)", "Description",
                           "Pièces changées (designation:qté:coût; ...)"],
                          [["2026-01-10", "Presse à injection AM-1", "Articles Ménagers", "Correctif", "Ousmane Diop",
                            3.5, 15000, "Remplacement courroie", "Courroie trapézoïdale:1:12000; Roulement 6205:2:4500"]]),
        "Planning": (["Titre", "Date (AAAA-MM-JJ)", "Machine", "Département (optionnel)", "Type", "Statut", "Description"],
                     [["Graissage mensuel", "2026-02-05", "Presse à injection AM-1", "Articles Ménagers",
                       "Lubrification", "Planifié", "Graissage des paliers"]]),
        "Instructions": (["ADMI — Modèle d'import Excel"],
                         [["Chaque feuille = un type de données. Ne renommez pas les feuilles."],
                          ["Si une machine citée dans Arrêts/Interventions/Planning n'existe pas,"],
                          ["renseignez la colonne « Département » : ADMI la crée automatiquement."],
                          [""],
                          ["Départements valides :"]] +
                         [[f"- {d['nom']} ({d['court']})"] for d in DEPARTEMENTS] +
                         [[""], ["Statuts machine :"]] + [[f"- {s}"] for s in STATUTS_MACHINE] +
                         [[""], ["Types d'arrêt :"]] + [[f"- {s}"] for s in TYPES_ARRET] +
                         [[""], ["Types d'intervention :"]] + [[f"- {s}"] for s in TYPES_INTERV] +
                         [[""], ["Types de planning :"]] + [[f"- {s}"] for s in TYPES_PLAN] +
                         [[""], ["Format pièces : designation:quantité:coût séparés par ';'"],
                          ["Exemple : Courroie:1:12000; Roulement:2:4500"]]),
    }
    return _write_sheets(sheets)
